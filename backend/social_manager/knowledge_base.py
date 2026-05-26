"""
Knowledge base module for managing documents, PDFs, and training data.
Supports PDF, DOCX, TXT, CSV parsing and storage.
"""

from __future__ import annotations
import os
import json
from datetime import datetime
from typing import List, Dict, Optional, Any
from pathlib import Path
import tempfile
import logging

# File handling imports
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import csv
except ImportError:
    csv = None

from sqlalchemy import Column, Integer, String, Text, DateTime, JSON
from sqlalchemy.orm import declarative_base, sessionmaker
from social_manager.db import Base, SessionLocal, engine

logger = logging.getLogger(__name__)

# ===== DATABASE TABLE =====

class KnowledgeDocument(Base):
    """Database model for knowledge base documents."""
    __tablename__ = "knowledge_documents"
    id = Column(Integer, primary_key=True)
    filename = Column(String, nullable=False, index=True)
    category = Column(String, nullable=False, index=True)  # e.g., brand_voice, social_strategy
    file_type = Column(String, nullable=False)  # pdf, docx, txt, csv
    content = Column(Text, nullable=False)  # Extracted text content
    doc_metadata = Column(JSON, default={})  # Custom metadata
    uploaded_at = Column(DateTime, default=datetime.utcnow, index=True)
    processing_status = Column(String, default="completed")  # pending, processing, completed, failed


# ===== FILE PARSERS =====

class DocumentParser:
    """Parses various document formats into plain text."""
    
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """Extract text from PDF file."""
        if not PyPDF2:
            logger.warning("PyPDF2 not installed; returning filename instead")
            return f"[PDF Document: {os.path.basename(file_path)}]"
        
        try:
            content = []
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        content.append(f"--- Page {page_num + 1} ---\n{text}")
            return "\n".join(content) if content else "[PDF: No text extracted]"
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            return f"[Error parsing PDF: {str(e)}]"
    
    @staticmethod
    def parse_docx(file_path: str) -> str:
        """Extract text from DOCX file."""
        if not Document:
            logger.warning("python-docx not installed; returning filename instead")
            return f"[DOCX Document: {os.path.basename(file_path)}]"
        
        try:
            doc = Document(file_path)
            content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            return "\n".join(content) if content else "[DOCX: No text extracted]"
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            return f"[Error parsing DOCX: {str(e)}]"
    
    @staticmethod
    def parse_txt(file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"TXT parsing error: {e}")
            return f"[Error reading TXT: {str(e)}]"
    
    @staticmethod
    def parse_csv(file_path: str) -> str:
        """Convert CSV to formatted text."""
        if not csv:
            return "[CSV file uploaded but csv module not available]"
        
        try:
            content = []
            with open(file_path, 'r', encoding='utf-8') as file:
                reader = csv.DictReader(file)
                for row_num, row in enumerate(reader, 1):
                    content.append(f"Row {row_num}: {json.dumps(row)}")
            return "\n".join(content) if content else "[CSV: No rows found]"
        except Exception as e:
            logger.error(f"CSV parsing error: {e}")
            return f"[Error reading CSV: {str(e)}]"
    
    @staticmethod
    def parse_file(file_path: str, file_type: str) -> str:
        """Parse any supported file type."""
        file_type = file_type.lower().strip('.')
        
        parsers = {
            'pdf': DocumentParser.parse_pdf,
            'docx': DocumentParser.parse_docx,
            'txt': DocumentParser.parse_txt,
            'csv': DocumentParser.parse_csv,
        }
        
        parser = parsers.get(file_type, DocumentParser.parse_txt)
        return parser(file_path)


# ===== KNOWLEDGE BASE MANAGER =====

class KnowledgeBaseManager:
    """Manages knowledge base documents: storage, retrieval, search."""
    
    def __init__(self, db_session=None):
        self.db_session = db_session or SessionLocal()
        self.parser = DocumentParser()
    
    def add_document(self, file_path: str, filename: str, category: str, 
                    file_type: str, metadata: Optional[Dict] = None) -> KnowledgeDocument:
        """
        Add a document to the knowledge base.
        
        Args:
            file_path: Path to the file
            filename: Display filename
            category: Document category (e.g., 'brand_voice', 'social_strategy')
            file_type: File extension (pdf, docx, txt, csv)
            metadata: Optional custom metadata dict
            
        Returns:
            KnowledgeDocument record
        """
        try:
            # Parse file content
            content = self.parser.parse_file(file_path, file_type)
            
            # Create database record
            doc = KnowledgeDocument(
                filename=filename,
                category=category,
                file_type=file_type.lower().strip('.'),
                content=content,
                doc_metadata=metadata or {},
                processing_status="completed"
            )
            
            self.db_session.add(doc)
            self.db_session.commit()
            
            logger.info(f"✓ Document added: {filename} (category: {category})")
            return doc
            
        except Exception as e:
            self.db_session.rollback()
            logger.error(f"✗ Failed to add document: {e}")
            raise
    
    def get_documents(self, category: Optional[str] = None) -> List[KnowledgeDocument]:
        """Get all documents, optionally filtered by category."""
        query = self.db_session.query(KnowledgeDocument)
        if category:
            query = query.filter_by(category=category)
        return query.order_by(KnowledgeDocument.uploaded_at.asc()).all()
    
    def get_document(self, doc_id: int) -> Optional[KnowledgeDocument]:
        """Get a specific document by ID."""
        return self.db_session.query(KnowledgeDocument).filter_by(id=doc_id).first()
    
    def delete_document(self, doc_id: int) -> bool:
        """Delete a document from the knowledge base."""
        try:
            doc = self.get_document(doc_id)
            if doc:
                self.db_session.delete(doc)
                self.db_session.commit()
                logger.info(f"✓ Document deleted: {doc.filename}")
                return True
            return False
        except Exception as e:
            self.db_session.rollback()
            logger.error(f"✗ Failed to delete document: {e}")
            return False
    
    def search_documents(self, query: str, category: Optional[str] = None) -> List[Dict]:
        """
        Search documents by content keyword.
        
        Args:
            query: Search term
            category: Optional category filter
            
        Returns:
            List of matching documents with relevance scores
        """
        docs = self.get_documents(category)
        results = []
        query_lower = query.lower()
        
        for doc in docs:
            # Simple keyword matching with relevance scoring
            matches = doc.content.lower().count(query_lower)
            if matches > 0:
                results.append({
                    "id": doc.id,
                    "filename": doc.filename,
                    "category": doc.category,
                    "type": doc.file_type,
                    "uploaded_at": doc.uploaded_at.isoformat(),
                    "relevance_score": min(matches / 10, 1.0),  # Normalized 0-1
                    "content_preview": doc.content[:200] + "..."
                })
        
        return sorted(results, key=lambda x: x['relevance_score'], reverse=True)
    
    def get_category_summary(self, category: str) -> Dict[str, Any]:
        """Get summary of documents in a category."""
        docs = self.get_documents(category)
        
        return {
            "category": category,
            "document_count": len(docs),
            "documents": [
                {
                    "id": doc.id,
                    "filename": doc.filename,
                    "type": doc.file_type,
                    "size_estimate": len(doc.content),
                    "uploaded_at": doc.uploaded_at.isoformat()
                }
                for doc in docs
            ]
        }
    
    def get_all_categories(self) -> List[str]:
        """Get all unique categories in the knowledge base."""
        results = self.db_session.query(KnowledgeDocument.category).distinct().all()
        return [r[0] for r in results]
    
    def build_context_for_llm(self, category: Optional[str] = None, max_chars: int = 8000) -> str:
        """
        Build context string from knowledge base for LLM prompting.
        
        Args:
            category: Optional category filter
            max_chars: Maximum characters to include
            
        Returns:
            Formatted knowledge base context
        """
        docs = self.get_documents(category)
        if not docs:
            return "[No documents in knowledge base]"
        
        context_parts = []
        total_chars = 0
        
        for doc in docs:
            # Truncate individual doc if needed
            truncated = doc.content[:2000]
            part = f"## {doc.filename}\n**Category:** {doc.category}\n\n{truncated}\n"
            
            if total_chars == 0 or total_chars + len(part) <= max_chars:
                context_parts.append(part)
                total_chars += len(part)
            else:
                break
        
        return "\n".join(context_parts) if context_parts else "[No documents available for context]"


# ===== INITIALIZATION =====

def init_knowledge_base_with_samples(db_session=None):
    """Initialize knowledge base with sample documents."""
    db_session = db_session or SessionLocal()
    
    # Create knowledge_documents table
    Base.metadata.create_all(bind=engine)
    
    manager = KnowledgeBaseManager(db_session)
    
    # Check if already initialized
    existing = db_session.query(KnowledgeDocument).count()
    if existing > 0:
        logger.info(f"✓ Knowledge base already populated with {existing} documents")
        return manager
    
    # Sample documents to load
    docs_dir = Path(__file__).parent.parent / "docs"
    samples = [
        {
            "path": docs_dir / "brand_voice_guidelines.txt",
            "filename": "brand_voice_guidelines.txt",
            "category": "Brand voice",
            "type": "txt"
        },
        {
            "path": docs_dir / "social_content_pillars_q2.txt",
            "filename": "social_content_pillars_q2.txt",
            "category": "Social strategy",
            "type": "txt"
        },
        {
            "path": docs_dir / "audience_segments.csv",
            "filename": "audience_segments.csv",
            "category": "Target audience",
            "type": "csv"
        },
        {
            "path": docs_dir / "competitor_social_audit.txt",
            "filename": "competitor_social_audit.txt",
            "category": "Competitors",
            "type": "txt"
        },
        {
            "path": docs_dir / "campaign_brief_spring.txt",
            "filename": "campaign_brief_spring.txt",
            "category": "Campaign briefs",
            "type": "txt"
        }
    ]
    
    for sample in samples:
        try:
            if sample["path"].exists():
                manager.add_document(
                    file_path=str(sample["path"]),
                    filename=sample["filename"],
                    category=sample["category"],
                    file_type=sample["type"]
                )
        except Exception as e:
            logger.warning(f"Could not load sample {sample['filename']}: {e}")
    
    return manager

def init_knowledge_base(seed: bool = False):
    """Initialize knowledge base tables and optionally seed demo data."""
    Base.metadata.create_all(bind=engine)
    
    if seed:
        # Add demo documents
        manager = KnowledgeBaseManager()
        demo_docs = [
            {
                "filename": "brand_voice_guidelines.txt",
                "category": "brand_voice",
                "file_type": "txt",
                "content": """BRAND VOICE GUIDELINES
Our brand voice is friendly, direct, and conversational.
- Be authentic and human
- Avoid corporate jargon
- Use active voice
- Keep sentences short and punchy
- Include emojis sparingly for warmth
- Focus on customer benefits, not features"""
            },
            {
                "filename": "social_strategy_framework.txt",
                "category": "social_strategy",
                "file_type": "txt",
                "content": """SOCIAL STRATEGY FRAMEWORK
Content Pillars:
1. Product education - 30%
2. Community engagement - 25%
3. Brand lifestyle - 25%
4. Sales/promotions - 20%

Posting frequency:
- Instagram: 3x per week
- LinkedIn: 2x per week
- TikTok: 2x per week
- Twitter/X: Daily"""
            },
            {
                "filename": "target_audience_personas.txt",
                "category": "audience",
                "file_type": "txt",
                "content": """TARGET AUDIENCE PERSONAS

Persona 1: Sarah (Age 28, Marketing Manager)
- Goals: Career growth, staying informed on trends
- Pain points: Time management, information overload
- Social platforms: LinkedIn, Instagram
- Content preference: Educational, practical tips

Persona 2: Alex (Age 35, Business Owner)
- Goals: Business growth, networking
- Pain points: Limited time, decision fatigue
- Social platforms: LinkedIn, Twitter
- Content preference: Industry insights, success stories"""
            }
        ]
        
        for demo in demo_docs:
            doc = KnowledgeDocument(
                filename=demo["filename"],
                category=demo["category"],
                file_type=demo["file_type"],
                content=demo["content"],
                processing_status="completed"
            )
            manager.db_session.add(doc)
        
        manager.db_session.commit()
        logger.info("✓ Knowledge base seeded with demo documents")


# ===== UTILITIES =====

def save_uploaded_file(file_content: bytes, filename: str, temp_dir: str = None) -> str:
    """
    Save uploaded file to temporary location.
    
    Args:
        file_content: File bytes
        filename: Original filename
        temp_dir: Optional temporary directory (creates one if not provided)
        
    Returns:
        Path to saved file
    """
    if not temp_dir:
        temp_dir = tempfile.gettempdir()
    
    file_path = os.path.join(temp_dir, filename)
    with open(file_path, 'wb') as f:
        f.write(file_content)
    
    return file_path


def get_file_extension(filename: str) -> str:
    """Extract and validate file extension."""
    ext = Path(filename).suffix.lower().strip('.')
    valid_extensions = ['pdf', 'docx', 'txt', 'csv']
    return ext if ext in valid_extensions else 'txt'


